import os
from pathlib import Path
from docx import Document

# Create a test DOCX file
doc = Document()
doc.add_paragraph("Test document content")
doc.save("test.docx")

# Verify the file was created
if not Path("test.docx").exists():
    print("Failed to create test.docx")
else:
    print("Created test.docx successfully")

# Try to read the file
try:
    doc_read = Document("test.docx")
    print("Successfully read test.docx")
    print(f"Content: {doc_read.paragraphs[0].text}")
except Exception as e:
    print(f"Error reading test.docx: {e}")
    import traceback
    traceback.print_exc()

